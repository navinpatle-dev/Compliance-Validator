from pathlib import Path
import fitz  # PyMuPDF
import docx
import logging

# --- Configuration ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file."""
    try:
        with fitz.open(file_path) as doc:
            text = "".join(page.get_text() for page in doc)
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a .docx file."""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        raise

def create_modified_docx(text: str, output_dir: Path, task_id: str) -> str:
    """Creates a new .docx file with the provided text."""
    try:
        doc = docx.Document()
        doc.add_paragraph(text)
        
        output_path = output_dir / f"modified_{task_id}.docx"
        doc.save(output_path)
        return str(output_path)
    except Exception as e:
        logger.error(f"Error creating modified DOCX file: {e}")
        raise