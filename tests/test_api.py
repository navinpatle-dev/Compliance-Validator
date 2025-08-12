import os
import time
from fastapi.testclient import TestClient
from pathlib import Path
import sys

# Add the parent directory to the sys.path to allow imports from the 'app' module
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.main import app

# --- Test Client Setup ---
client = TestClient(app)

# --- Test Data ---
TEST_FILES_DIR = Path(__file__).parent / "sample_docs"
COMPLIANT_DOC_PATH = TEST_FILES_DIR / "compliant_doc.docx"
NON_COMPLIANT_DOC_PATH = TEST_FILES_DIR / "non_compliant_doc.docx"
INVALID_FILE_PATH = TEST_FILES_DIR / "invalid_file.txt" # Create this file for testing

def setup_module(module):
    """Create dummy files for testing."""
    TEST_FILES_DIR.mkdir(exist_ok=True)
    
    # Create a dummy compliant docx
    try:
        import docx
        doc = docx.Document()
        doc.add_paragraph("This is a well-written document. It follows all guidelines and is perfectly clear.")
        doc.save(COMPLIANT_DOC_PATH)
        
        # Create a dummy non-compliant docx
        doc = docx.Document()
        doc.add_paragraph("This document have many error. It's clarity is bad. The rules was broken by the writer.")
        doc.save(NON_COMPLIANT_DOC_PATH)

        # Create an invalid file
        with open(INVALID_FILE_PATH, "w") as f:
            f.write("This is a text file.")
            
    except ImportError:
        print("python-docx is not installed. Skipping creation of test docx files.")


def test_read_root():
    """Test the root endpoint."""
    response = client.get("/")
    assert response.status_code == 200
    assert "LangChain/Gemini Edition" in response.json()["message"]

def test_upload_invalid_file_type():
    """Test uploading a file with an unsupported extension."""
    with open(INVALID_FILE_PATH, "rb") as f:
        response = client.post("/check-compliance/", files={"file": ("invalid_file.txt", f, "text/plain")})
    assert response.status_code == 400
    assert "Invalid file type" in response.json()["detail"]

def test_compliance_check_workflow():
    """Test the full workflow: upload, poll for results."""
    # Ensure the API key is set for this test
    if not os.getenv("GEMINI_API_KEY"):
        print("Skipping workflow test: GEMINI_API_KEY not set.")
        return

    with open(NON_COMPLIANT_DOC_PATH, "rb") as f:
        response = client.post(
            "/check-compliance/",
            files={"file": ("non_compliant_doc.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
    
    assert response.status_code == 202
    data = response.json()
    assert "task_id" in data
    task_id = data["task_id"]

    # Poll for results
    timeout = 60  # seconds
    start_time = time.time()
    while time.time() - start_time < timeout:
        res_response = client.get(f"/results/{task_id}")
        assert res_response.status_code == 200
        res_data = res_response.json()
        if res_data["status"] == "completed":
            assert res_data["report"]["summary"]["compliance_status"] == "Non-Compliant"
            assert len(res_data["report"]["violations"]) > 0
            return
        elif res_data["status"] == "failed":
            assert False, f"Task failed with report: {res_data.get('report')}"
        time.sleep(2)

    assert False, "Polling for results timed out."

def test_get_results_not_found():
    """Test retrieving results for a non-existent task."""
    response = client.get("/results/non-existent-task-id")
    assert response.status_code == 404